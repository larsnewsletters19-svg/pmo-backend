"""
AI PMO Generator - Backend Server
Flask server that generates formatted Word documents using python-docx
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import requests
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from io import BytesIO
import tempfile

app = Flask(__name__)

# TEMPORARY: Allow all origins for testing
# This will help us debug if CORS works at all
CORS(app, origins="*", supports_credentials=False)

print("=" * 60)
print("🔓 CORS: Allowing ALL origins (temporary for testing)")
print("=" * 60)


def create_formatted_docx(markdown_text):
    """Convert markdown text to formatted Word document"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)
    
    lines = markdown_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # H1 - Rubrik1
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_paragraph(text)
            p.style = doc.styles['Heading 1']
            run = p.runs[0]
            run.font.name = 'Segoe UI Semibold'
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1A, 0x4D, 0x80)
            run.bold = True
        
        # H2 - Rubrik2
        elif line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_paragraph(text)
            p.style = doc.styles['Heading 2']
            run = p.runs[0]
            run.font.name = 'Segoe UI Semibold'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x41, 0x6D, 0x94)
            run.bold = True
        
        # H3 - Rubrik3
        elif line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph(text)
            p.style = doc.styles['Heading 3']
            run = p.runs[0]
            run.font.name = 'Segoe UI Semibold'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            run.bold = True
        
        # Bullet point
        elif line.startswith('* ') or line.startswith('• '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            run = p.runs[0]
            run.font.name = 'Segoe UI'
            run.font.size = Pt(11)
        
        elif line.startswith('- '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            run = p.runs[0]
            run.font.name = 'Segoe UI'
            run.font.size = Pt(11)
        
        # Numbered list
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            p = doc.add_paragraph(text, style='List Number')
            run = p.runs[0]
            run.font.name = 'Segoe UI'
            run.font.size = Pt(11)
        
        # Table detection (markdown table)
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # Parse table
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            i -= 1  # Back up one line
            
            # Filter out separator lines
            table_rows = [l for l in table_lines if not re.match(r'^[\|\s\-:]+$', l)]
            
            if len(table_rows) > 0:
                # Parse columns
                cols = [cell.strip() for cell in table_rows[0].split('|') if cell.strip()]
                num_cols = len(cols)
                num_rows = len(table_rows)
                
                # Create table
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_text in enumerate(table_rows):
                    cells = [cell.strip() for cell in row_text.split('|') if cell.strip()]
                    for col_idx, cell_text in enumerate(cells[:num_cols]):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_text
                        
                        # Format cell
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Segoe UI'
                                run.font.size = Pt(10)
                                
                                # Header row - bold
                                if row_idx == 0:
                                    run.bold = True
                
                doc.add_paragraph()  # Add spacing after table
        
        # Normal paragraph
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = 'Segoe UI'
                run.font.size = Pt(11)
        
        i += 1
    
    return doc

@app.route('/api/generate', methods=['OPTIONS'])
def generate_options():
    """Handle preflight request explicitly"""
    response = jsonify({'status': 'ok'})
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,POST,OPTIONS')
    return response, 200

@app.route('/api/generate', methods=['POST'])
def generate():
    """Handle generation requests and return both text and .docx file"""
    try:
        data = request.json
        # Try to get API key from request, fallback to environment variable
        api_key = data.get('api_key') or os.environ.get('CLAUDE_API_KEY')
        prompt = data.get('prompt')
        
        if not api_key:
            return jsonify({'error': 'Missing API key. Please configure CLAUDE_API_KEY in Railway or provide api_key in request'}), 400
        
        if not prompt:
            return jsonify({'error': 'Missing prompt'}), 400
        
        # Call Claude API
        response = requests.post(
            'https://api.anthropic.com/v1/messages',
            headers={
                'Content-Type': 'application/json',
                'x-api-key': api_key,
                'anthropic-version': '2023-06-01'
            },
            json={
                'model': 'claude-sonnet-4-20250514',
                'max_tokens': 4000,
                'messages': [
                    {
                        'role': 'user',
                        'content': prompt
                    }
                ]
            }
        )
        
        if response.status_code != 200:
            return jsonify({'error': response.json()}), response.status_code
        
        claude_response = response.json()
        content = claude_response['content'][0]['text']
        
        # AGGRESSIVE CLEANING: Force single version
        
        # Step 1: Remove all version headers
        cleaned = re.sub(r'^[\s\n]*#{1,2}\s*(?:OneNote|Word)[- ]?version\s*:?\s*\n+', '', content, flags=re.IGNORECASE | re.MULTILINE)
        cleaned = re.sub(r'\n+[\s]*#{1,2}\s*(?:OneNote|Word)[- ]?version\s*:?\s*\n+', '\n\n', cleaned, flags=re.IGNORECASE)
        
        # Step 2: Remove all separator lines
        cleaned = re.sub(r'\n\s*[-=]{3,}\s*\n', '\n\n', cleaned)
        
        # Step 3: Intelligent duplication detection
        # Strip ALL formatting and compare raw content
        def strip_formatting(text):
            """Remove all markdown formatting to get raw content"""
            stripped = re.sub(r'#{1,6}\s+', '', text)  # Remove headers
            stripped = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)  # Remove bold
            stripped = re.sub(r'\*(.+?)\*', r'\1', stripped)  # Remove italic
            stripped = re.sub(r'\|', '', stripped)  # Remove table pipes
            stripped = re.sub(r'[-:]+', '', stripped)  # Remove table separators
            stripped = re.sub(r'\s+', ' ', stripped)  # Normalize whitespace
            return stripped.strip().lower()
        
        # Split content into paragraphs
        paragraphs = [p.strip() for p in cleaned.split('\n\n') if p.strip()]
        
        if len(paragraphs) > 5:  # Only check substantial content
            # Compare first half with second half (stripped of formatting)
            midpoint = len(paragraphs) // 2
            
            first_half_stripped = ' '.join([strip_formatting(p) for p in paragraphs[:midpoint]])
            second_half_stripped = ' '.join([strip_formatting(p) for p in paragraphs[midpoint:]])
            
            # Calculate similarity (how much of first half appears in second half)
            first_words = set(first_half_stripped.split())
            second_words = set(second_half_stripped.split())
            
            if first_words and second_words:
                overlap = len(first_words & second_words)
                overlap_ratio = overlap / len(first_words)
                
                # If >75% word overlap, it's duplicated content with different formatting
                if overlap_ratio > 0.75:
                    # Take first half only (keep original formatting)
                    cleaned = '\n\n'.join(paragraphs[:midpoint])
        
        # Step 4: Clean up excessive whitespace
        cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
        cleaned = cleaned.strip()
        
        # Use cleaned content for Word version
        word_text = cleaned
        
        # Create simplified OneNote version
        onenote_version = re.sub(r'#{1,3}\s+', '', cleaned)
        onenote_version = re.sub(r'\*\*(.+?)\*\*', r'\1', onenote_version)
        onenote_version = re.sub(r'\|[^\n]+\|', '', onenote_version)
        onenote_version = re.sub(r'\n\s*[-:]+\s*\n', '\n', onenote_version)
        onenote_version = re.sub(r'\n{3,}', '\n\n', onenote_version).strip()
        
        # Generate .docx file
        doc = create_formatted_docx(word_text)
        
        # Create persistent directory for generated files
        output_dir = os.path.join(os.path.dirname(__file__), 'generated_docs')
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate unique filename
        import time
        filename = f'pmo_doc_{int(time.time() * 1000)}.docx'
        filepath = os.path.join(output_dir, filename)
        
        # Save document
        doc.save(filepath)
        
        # Return response with filename (not full path for security)
        return jsonify({
            'content': claude_response['content'],
            'onenote_version': onenote_version,
            'word_version': word_text,
            'docx_file': filename
        })
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/download/<path:filename>', methods=['OPTIONS'])
def download_options(filename):
    """Handle preflight for download"""
    response = jsonify({'status': 'ok'})
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,OPTIONS')
    return response, 200

@app.route('/api/download/<path:filename>', methods=['GET'])
def download(filename):
    """Download generated .docx file"""
    try:
        # Security: only allow filenames, no path traversal
        if '/' in filename or '\\' in filename:
            return jsonify({'error': 'Invalid filename'}), 400
        
        # Get file from generated_docs directory
        output_dir = os.path.join(os.path.dirname(__file__), 'generated_docs')
        filepath = os.path.join(output_dir, filename)
        
        if not os.path.exists(filepath):
            return jsonify({'error': 'File not found'}), 404
        
        return send_file(
            filepath,
            as_attachment=True,
            download_name='pmo_document.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint"""
    return jsonify({'status': 'ok'})

if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 5000))
    host = os.environ.get('HOST', '0.0.0.0')
    
    print("=" * 60)
    print("🚀 AI PMO Generator Backend Server v2.0")
    print("=" * 60)
    print("✨ NY FUNKTION: Genererar riktiga .docx-filer!")
    print(f"Server körs på: http://{host}:{port}")
    print("Öppna nu ai-pmo-generator.html i din webbläsare")
    print("=" * 60)
    
    # Production mode if PORT is set by hosting platform
    debug_mode = os.environ.get('PORT') is None
    app.run(debug=debug_mode, port=port, host=host)
